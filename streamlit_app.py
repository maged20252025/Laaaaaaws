import streamlit as st
import streamlit.components.v1 as components
from docx import Document
import re
import os
import time
import base64

st.set_page_config(page_title="القوانين اليمنية بآخر تعديلاتها حتى عام 2025م", layout="wide")
st.markdown("<h1 style='text-align: center;'>مرحبًا بك في تطبيق القوانين اليمنية بآخر تعديلاتها حتى عام 2025م</h1>", unsafe_allow_html=True)

def is_activated():
    return os.path.exists("activated.txt")

def activate_app(code):
    if not os.path.exists("activation_codes.txt"):
        return False
    with open("activation_codes.txt", "r") as f:
        codes = [line.strip() for line in f.readlines()]
    if code in codes:
        codes.remove(code)
        with open("activation_codes.txt", "w") as f:
            for c in codes:
                f.write(c + "\n")
        with open("activated.txt", "w") as f:
            f.write("activated")
        return True
    return False

def highlight_keywords(text, keywords):
    # التأكد من أن النص هو من نوع string
    text = str(text)
    for kw in keywords:
        # استخدام re.escape للتأكد من التعامل الصحيح مع المحارف الخاصة في الكلمات المفتاحية
        # استخدام re.IGNORECASE للبحث غير الحساس لحالة الأحرف
        # استخدام re.UNICODE لضمان التعامل الصحيح مع محارف Unicode (مثل العربية)
        text = re.sub(f"({re.escape(kw)})", r"<mark>\1</mark>", text, flags=re.IGNORECASE | re.UNICODE)
    return text

def extract_context(paragraphs, keywords, context_lines=5):
    # التأكد من أن المدخلات هي قائمة من النصوص
    paragraphs = [str(p) for p in paragraphs]
    
    # تحويل الكلمات المفتاحية إلى صيغة قابلة للبحث بـ regex
    # استخدام re.UNICODE هنا أيضًا
    search_pattern = re.compile('|'.join([re.escape(kw) for kw in keywords]), re.IGNORECASE | re.UNICODE)
    
    matched_indexes = []
    for i, line in enumerate(paragraphs):
        if search_pattern.search(line):
            matched_indexes.append(i)
            
    context_set = set()
    for idx in matched_indexes:
        # هنا سنعدل لزيادة السياق إلى 5 أسطر قبل و 5 أسطر بعد
        for i in range(max(0, idx - context_lines), min(len(paragraphs), idx + context_lines + 1)):
            context_set.add(i)
            
    # تصفية الفقرات الفارغة أو المسافات البيضاء
    filtered_paragraphs = [paragraphs[i] for i in sorted(context_set) if paragraphs[i].strip()]
    return "\n".join(filtered_paragraphs)


def export_results_to_docx(results, filename="نتائج_البحث.docx"):
    doc = Document()
    doc.add_heading("نتائج البحث", 0)
    for r in results:
        doc.add_heading(f'{r["law"]} - المادة {r["num"]}', level=1)
        doc.add_paragraph(r["context"]) # يمكننا استخدام 'context' هنا لأنه يمثل النص المستخرج
    # التأكد من وجود المجلد '/mnt/data/' إذا كنت تستخدم بيئة Docker (مثل Streamlit Cloud)
    # وإلا استخدم os.path.join(os.getcwd(), filename) لحفظه في نفس مجلد التطبيق
    # للبيئة المحلية، يمكن استخدام:
    filepath = os.path.join(os.getcwd(), filename) # سيتم حفظه في نفس مجلد تشغيل التطبيق
    doc.save(filepath)
    return filepath

def run_main_app():
    components.html("""
    <style>
    .scroll-btn {
        position: fixed;
        left: 10px;
        padding: 12px;
        font-size: 24px;
        border-radius: 50%;
        background-color: #c5e1a5;
        color: black;
        cursor: pointer;
        z-index: 9999;
        border: none;
        box-shadow: 1px 1px 5px #888;
    }
    #scroll-top-btn { bottom: 80px; }
    #scroll-bottom-btn { bottom: 20px; }
    </style>
    <button class='scroll-btn' id='scroll-top-btn' onclick='window.scrollTo({top: 0, behavior: "smooth"});'>⬆️</button>
    <button class='scroll-btn' id='scroll-bottom-btn' onclick='window.scrollTo({top: document.body.scrollHeight, behavior: "smooth"});'>⬇️</button>
    """, height=1)

    subfolders = [f.path for f in os.scandir() if f.is_dir() and f.name not in [".git", ".streamlit"]]
    if not subfolders:
        st.warning("📂 لا توجد مجلدات قوانين.")
        return

    selected_folder = st.selectbox("اختر مجلدًا للبحث فيه:", ["🔍 كل المجلدات"] + subfolders)

    all_files = {}
    if selected_folder == "🔍 كل المجلدات":
        for folder in subfolders:
            files = [f for f in os.listdir(folder) if f.endswith(".docx")]
            all_files[folder] = files
    else:
        files = [f for f in os.listdir(selected_folder) if f.endswith(".docx")]
        all_files[selected_folder] = files

    keywords = st.text_area("الكلمات المفتاحية (افصل بفاصلة)", "")

    if "results" not in st.session_state:
        st.session_state.results = []
    if "search_done" not in st.session_state:
        st.session_state.search_done = False

    if st.button("🔍 بدء البحث") and keywords:
        kw_list = [k.strip() for k in keywords.split(",") if k.strip()]
        results = []

        for folder, files in all_files.items():
            for file in files:
                doc_path = os.path.join(folder, file)
                try:
                    doc = Document(doc_path)
                except Exception as e:
                    st.warning(f"⚠️ تعذر قراءة الملف {file} في المجلد {folder}: {e}. قد يكون الملف تالفًا أو مشفرًا.")
                    continue

                law_name = file.replace(".docx", "")
                
                # تخزين جميع فقرات الوثيقة لمعالجتها كوحدة واحدة للبحث عن المادة
                all_paragraphs_in_doc = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
                
                # تحديد المادة الحالية والنص الخاص بها
                current_article_paragraphs = []
                last_article_num = "غير معروفة"

                # مراجعة جميع الفقرات في المستند
                for i, para_text in enumerate(all_paragraphs_in_doc):
                    match = re.match(r"مادة\s*\(?\s*(\d+)\)?", para_text)
                    if match:
                        # إذا وجدنا مادة جديدة
                        if current_article_paragraphs:
                            # قبل البدء في المادة الجديدة، نقوم بمعالجة المادة السابقة
                            full_article_text = "\n".join(current_article_paragraphs)
                            # البحث عن الكلمات المفتاحية في كامل نص المادة
                            if any(kw.lower() in full_article_text.lower() for kw in kw_list):
                                # استخراج السياق من كل فقرات المادة الحالية
                                # تمرير الفقرات الأصلية وليس full_article_text للحفاظ على الفواصل
                                context = extract_context(current_article_paragraphs, kw_list, context_lines=5) # تم تعديل هنا
                                results.append({
                                    "law": law_name,
                                    "num": last_article_num,
                                    "text": highlight_keywords(context, kw_list),
                                    "plain": full_article_text,
                                    "context": context,
                                    "keywords": kw_list
                                })
                            current_article_paragraphs = [] # تفريغ لمادة جديدة
                        last_article_num = match.group(1) # تحديث رقم المادة
                        current_article_paragraphs.append(para_text) # إضافة سطر المادة نفسه
                    else:
                        current_article_paragraphs.append(para_text)
                
                # معالجة آخر مادة في المستند بعد انتهاء الحلقة
                if current_article_paragraphs:
                    full_article_text = "\n".join(current_article_paragraphs)
                    if any(kw.lower() in full_article_text.lower() for kw in kw_list):
                        context = extract_context(current_article_paragraphs, kw_list, context_lines=5) # تم تعديل هنا
                        results.append({
                            "law": law_name,
                            "num": last_article_num,
                            "text": highlight_keywords(context, kw_list),
                            "plain": full_article_text,
                            "context": context,
                            "keywords": kw_list
                        })

        st.session_state.results = results
        st.session_state.search_done = True

    if st.session_state.search_done and st.session_state.results:
        results = st.session_state.results
        unique_laws = sorted(set(r["law"] for r in results))
        st.success(f"تم العثور على {len(results)} نتيجة في {len(unique_laws)} قانون/ملف.")
        
        # إضافة خيار 'الكل' في الفلترة حسب القانون
        selected_law = st.selectbox("فلترة حسب القانون", ["الكل"] + unique_laws)
        filtered = results if selected_law == "الكل" else [r for r in results if r["law"] == selected_law]

        # عرض النتائج
        for r in filtered:
            st.markdown(f"""
<div style="background-color:#f1f8e9;padding:15px;margin-bottom:15px;border-radius:10px;
            border:1px solid #c5e1a5;direction:rtl;text-align:right; overflow-wrap: break-word;">
    <p style="font-weight:bold;font-size:18px;margin:0">🔷 {r["law"]} - المادة {r["num"]}</p>
    <p style="font-size:17px;line-height:1.8;margin-top:10px">
        {r["text"]}
    </p>
</div>
""", unsafe_allow_html=True)

        if filtered:
            filepath = export_results_to_docx(filtered)
            with open(filepath, "rb") as f:
                st.download_button(
                    label="📥 تحميل النتائج كملف Word",
                    data=f,
                    file_name="نتائج_البحث.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

def main():
    if not is_activated():
        st.warning("⚠️ التطبيق غير مفعل. يرجى التفعيل أو استخدام النسخة التجريبية.")
        col1, col2 = st.columns(2)
        with col1:
            code = st.text_input("أدخل كود التفعيل هنا", key="activation_code_input") # إضافة مفتاح فريد
            if st.button("🔐 تفعيل التطبيق", key="activate_button"): # إضافة مفتاح فريد
                if code and activate_app(code.strip()):
                    st.success("✅ تم التفعيل بنجاح! يرجى إعادة تشغيل التطبيق لتطبيق التغييرات.")
                    st.session_state.activated = True # تحديث حالة التفعيل في الجلسة
                    st.experimental_rerun() # إعادة تشغيل التطبيق
                else:
                    st.error("❌ كود التفعيل غير صحيح أو انتهى.")
        with col2:
            if "trial_start" not in st.session_state:
                if st.button("🕒 بدء التجربة المجانية", key="start_trial_button"): # إضافة مفتاح فريد
                    st.session_state.trial_start = time.time()
                    st.success("🎉 بدأت النسخة التجريبية. لديك ساعة واحدة.")
                    st.experimental_rerun() # إعادة تشغيل لعرض التطبيق
            elif time.time() - st.session_state.trial_start < 3600:
                st.info(f"✅ النسخة التجريبية نشطة. تبقى لديك حوالي {int((3600 - (time.time() - st.session_state.trial_start)) / 60)} دقيقة.")
                run_main_app()
            else:
                st.error("❌ انتهت مدة التجربة المجانية. يرجى التفعيل.")
    else:
        run_main_app()

main()
